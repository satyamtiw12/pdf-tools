from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
import uvicorn
import os, shutil, uuid, zipfile
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional, List
import threading, time, io

from pypdf import PdfReader, PdfWriter
from PIL import Image
import img2pdf
import fitz
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib import colors as rl_colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from pdf2docx import Converter as PdfToDocxConverter
from docx import Document
import pdfplumber
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

app = FastAPI(title="PDFBox API", version="1.1.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

BASE_DIR   = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
FRONTEND   = BASE_DIR.parent / "index.html"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

MAX_FILE_MB = 20
AUTO_DELETE_H = 2

def save_upload(file: UploadFile) -> Path:
    ext  = Path(file.filename).suffix or ".bin"
    dest = UPLOAD_DIR / f"{uuid.uuid4().hex}{ext}"
    with open(dest, "wb") as f: shutil.copyfileobj(file.file, f)
    if dest.stat().st_size / 1024**2 > MAX_FILE_MB:
        dest.unlink(); raise HTTPException(400, f"File too large. Max {MAX_FILE_MB} MB.")
    return dest

def out_path(name): return OUTPUT_DIR / f"{uuid.uuid4().hex}_{name}"

def parse_page_range(expr, total):
    pages = set()
    for part in expr.split(","):
        part = part.strip()
        if "-" in part:
            a,b = part.split("-",1)
            try: pages.update(range(int(a)-1, int(b)))
            except: pass
        elif part.isdigit(): pages.add(int(part)-1)
    return sorted(p for p in pages if 0 <= p < total)

def auto_cleaner():
    while True:
        cutoff = datetime.now() - timedelta(hours=AUTO_DELETE_H)
        for d in [UPLOAD_DIR, OUTPUT_DIR]:
            for f in d.iterdir():
                if f.is_file() and datetime.fromtimestamp(f.stat().st_mtime) < cutoff:
                    f.unlink(missing_ok=True)
        time.sleep(600)

threading.Thread(target=auto_cleaner, daemon=True).start()

@app.get("/")
async def root():
    return FileResponse(FRONTEND) if FRONTEND.exists() else {"message": "PDFBox API running!"}

@app.get("/health")
async def health(): return {"status": "ok", "tools": 22}

@app.post("/api/merge")
async def merge_pdf(files: List[UploadFile] = File(...)):
    if len(files) < 2: raise HTTPException(400, "At least 2 PDF files required.")
    paths = [save_upload(f) for f in files]
    writer = PdfWriter()
    for p in paths:
        for page in PdfReader(str(p)).pages: writer.add_page(page)
    out = out_path("merged.pdf")
    with open(out,"wb") as f: writer.write(f)
    for p in paths: p.unlink(missing_ok=True)
    return FileResponse(out, filename="merged.pdf", media_type="application/pdf")

@app.post("/api/split")
async def split_pdf(file: UploadFile=File(...), method: str=Form("each"), page_range: str=Form("")):
    path = save_upload(file); reader = PdfReader(str(path)); total = len(reader.pages)
    zip_out = out_path("split.zip")
    with zipfile.ZipFile(zip_out,"w") as zf:
        if method=="each" or not page_range.strip():
            for i,page in enumerate(reader.pages):
                w=PdfWriter(); w.add_page(page); buf=io.BytesIO(); w.write(buf)
                zf.writestr(f"page_{i+1}.pdf", buf.getvalue())
        else:
            indices=parse_page_range(page_range,total); w=PdfWriter()
            for i in indices: w.add_page(reader.pages[i])
            buf=io.BytesIO(); w.write(buf); zf.writestr("selected.pdf",buf.getvalue())
    path.unlink(missing_ok=True)
    return FileResponse(zip_out, filename="split_pages.zip", media_type="application/zip")

@app.post("/api/compress")
async def compress_pdf(file: UploadFile=File(...), level: str=Form("medium")):
    path=save_upload(file); orig=path.stat().st_size
    doc=fitz.open(str(path)); dpi={"low":150,"medium":100,"high":60}.get(level,100)
    writer=fitz.open()
    for page in doc:
        mat=fitz.Matrix(dpi/72,dpi/72); pix=page.get_pixmap(matrix=mat)
        np=writer.new_page(width=page.rect.width,height=page.rect.height)
        np.insert_image(np.rect,pixmap=pix)
    out=out_path("compressed.pdf"); writer.save(str(out),deflate=True,garbage=4)
    doc.close(); path.unlink(missing_ok=True)
    saved=round((1-out.stat().st_size/orig)*100) if orig>0 else 0
    return FileResponse(out, filename="compressed.pdf", media_type="application/pdf",
                        headers={"X-Saved-Percent":str(saved)})

@app.post("/api/pdf-to-jpg")
async def pdf_to_jpg(file: UploadFile=File(...), quality: str=Form("medium")):
    path=save_upload(file); doc=fitz.open(str(path))
    dpi={"low":72,"medium":150,"high":300}.get(quality,150)
    mat=fitz.Matrix(dpi/72,dpi/72); zip_out=out_path("pdf_images.zip")
    with zipfile.ZipFile(zip_out,"w",zipfile.ZIP_DEFLATED) as zf:
        for i,page in enumerate(doc): zf.writestr(f"page_{i+1}.jpg",page.get_pixmap(matrix=mat).tobytes("jpeg"))
    doc.close(); path.unlink(missing_ok=True)
    return FileResponse(zip_out, filename="pdf_images.zip", media_type="application/zip")

@app.post("/api/jpg-to-pdf")
async def jpg_to_pdf(files: List[UploadFile]=File(...), page_size: str=Form("a4")):
    paths=[save_upload(f) for f in files]; out=out_path("images.pdf")
    if page_size=="auto":
        with open(out,"wb") as f: f.write(img2pdf.convert([str(p) for p in paths]))
    else:
        size=A4 if page_size=="a4" else letter
        c=rl_canvas.Canvas(str(out),pagesize=size); W,H=size
        for p in paths:
            img=Image.open(p).convert("RGB"); iw,ih=img.size
            ratio=min(W/iw,H/ih); nw,nh=iw*ratio,ih*ratio
            c.drawImage(str(p),(W-nw)/2,(H-nh)/2,nw,nh); c.showPage()
        c.save()
    for p in paths: p.unlink(missing_ok=True)
    return FileResponse(out, filename="converted.pdf", media_type="application/pdf")

@app.post("/api/add-text")
async def add_text(file: UploadFile=File(...), text: str=Form(...), page_num: int=Form(1),
                   position: str=Form("bottom-center"), font_size: int=Form(14), color: str=Form("#000000")):
    path=save_upload(file); doc=fitz.open(str(path))
    page=doc[max(0,min(page_num-1,len(doc)-1))]; W,H=page.rect.width,page.rect.height
    pos={"top-left":(40,40),"top-center":(W/2,40),"top-right":(W-40,40),"center":(W/2,H/2),
         "bottom-left":(40,H-40),"bottom-center":(W/2,H-40),"bottom-right":(W-40,H-40)}
    x,y=pos.get(position,(W/2,H-40)); h=color.lstrip("#")
    r,g,b=tuple(int(h[i:i+2],16)/255 for i in (0,2,4))
    page.insert_text((x,y),text,fontsize=font_size,color=(r,g,b),overlay=True)
    out=out_path("text_added.pdf"); doc.save(str(out)); doc.close(); path.unlink(missing_ok=True)
    return FileResponse(out, filename="text_added.pdf", media_type="application/pdf")

@app.post("/api/add-signature")
async def add_signature(file: UploadFile=File(...), sig_text: str=Form(""),
                        page_num: int=Form(-1), sig_image: Optional[UploadFile]=File(None)):
    path=save_upload(file); doc=fitz.open(str(path))
    pi=len(doc)-1 if page_num==-1 else max(0,page_num-1)
    page=doc[pi]; W,H=page.rect.width,page.rect.height
    if sig_image and sig_image.filename:
        ip=save_upload(sig_image); page.insert_image(fitz.Rect(W-200,H-80,W-20,H-20),filename=str(ip)); ip.unlink(missing_ok=True)
    elif sig_text: page.insert_text((W-200,H-30),sig_text,fontsize=18,color=(0.1,0.1,0.6),overlay=True)
    out=out_path("signed.pdf"); doc.save(str(out)); doc.close(); path.unlink(missing_ok=True)
    return FileResponse(out, filename="signed.pdf", media_type="application/pdf")

@app.post("/api/highlight")
async def highlight_text(file: UploadFile=File(...), search_text: str=Form(...), color: str=Form("yellow")):
    rgb={"yellow":(1,1,0),"green":(0,1,0.5),"blue":(0.5,0.8,1),"pink":(1,0.5,0.8),"orange":(1,0.7,0.2)}.get(color,(1,1,0))
    path=save_upload(file); doc=fitz.open(str(path))
    for page in doc:
        for rect in page.search_for(search_text):
            a=page.add_highlight_annot(rect); a.set_colors(stroke=rgb); a.update()
    out=out_path("highlighted.pdf"); doc.save(str(out)); doc.close(); path.unlink(missing_ok=True)
    return FileResponse(out, filename="highlighted.pdf", media_type="application/pdf")

@app.post("/api/rotate")
async def rotate_pdf(file: UploadFile=File(...), degrees: int=Form(90), apply_to: str=Form("all")):
    path=save_upload(file); reader=PdfReader(str(path)); writer=PdfWriter()
    for i,page in enumerate(reader.pages):
        if apply_to=="all" or (apply_to=="first" and i==0): page.rotate(degrees)
        writer.add_page(page)
    out=out_path("rotated.pdf")
    with open(out,"wb") as f: writer.write(f)
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="rotated.pdf", media_type="application/pdf")

@app.post("/api/delete-pages")
async def delete_pages(file: UploadFile=File(...), pages_to_delete: str=Form(...)):
    path=save_upload(file); reader=PdfReader(str(path))
    to_del=set(parse_page_range(pages_to_delete,len(reader.pages))); writer=PdfWriter()
    for i,page in enumerate(reader.pages):
        if i not in to_del: writer.add_page(page)
    if not writer.pages: raise HTTPException(400,"At least one page must remain!")
    out=out_path("pages_deleted.pdf")
    with open(out,"wb") as f: writer.write(f)
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="pages_deleted.pdf", media_type="application/pdf")

@app.post("/api/reorder")
async def reorder_pages(file: UploadFile=File(...), new_order: str=Form(...)):
    path=save_upload(file); reader=PdfReader(str(path)); total=len(reader.pages)
    indices=[int(x.strip())-1 for x in new_order.split(",") if x.strip().isdigit()]
    indices=[i for i in indices if 0<=i<total]
    if not indices: raise HTTPException(400,"Enter valid page numbers.")
    writer=PdfWriter()
    for i in indices: writer.add_page(reader.pages[i])
    out=out_path("reordered.pdf")
    with open(out,"wb") as f: writer.write(f)
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="reordered.pdf", media_type="application/pdf")

@app.post("/api/protect")
async def protect_pdf(file: UploadFile=File(...), password: str=Form(...)):
    if len(password)<4: raise HTTPException(400,"Password must be at least 4 characters.")
    path=save_upload(file); reader=PdfReader(str(path)); writer=PdfWriter()
    for page in reader.pages: writer.add_page(page)
    writer.encrypt(password); out=out_path("protected.pdf")
    with open(out,"wb") as f: writer.write(f)
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="protected.pdf", media_type="application/pdf")

@app.post("/api/unlock")
async def unlock_pdf(file: UploadFile=File(...), password: str=Form(...)):
    path=save_upload(file); reader=PdfReader(str(path))
    if reader.is_encrypted:
        if not reader.decrypt(password): path.unlink(missing_ok=True); raise HTTPException(400,"Incorrect password!")
    writer=PdfWriter()
    for page in reader.pages: writer.add_page(page)
    out=out_path("unlocked.pdf")
    with open(out,"wb") as f: writer.write(f)
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="unlocked.pdf", media_type="application/pdf")

@app.post("/api/info")
async def pdf_info(file: UploadFile=File(...)):
    path=save_upload(file); reader=PdfReader(str(path)); info=reader.metadata or {}
    result={"pages":len(reader.pages),"file_size_kb":path.stat().st_size//1024,
            "title":str(info.get("/Title","N/A")),"author":str(info.get("/Author","N/A")),
            "encrypted":reader.is_encrypted}
    path.unlink(missing_ok=True); return JSONResponse(result)

@app.post("/api/pdf-to-word")
async def pdf_to_word(file: UploadFile=File(...)):
    path=save_upload(file); out=out_path("converted.docx")
    try:
        cv=PdfToDocxConverter(str(path)); cv.convert(str(out),start=0,end=None); cv.close()
    except Exception as e:
        path.unlink(missing_ok=True); raise HTTPException(500,f"Conversion failed: {e}")
    path.unlink(missing_ok=True)
    if not out.exists() or out.stat().st_size==0: raise HTTPException(500,"Output file is empty.")
    return FileResponse(out, filename="converted.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/api/word-to-pdf")
async def word_to_pdf(file: UploadFile=File(...)):
    path=save_upload(file); out=out_path("converted.pdf")
    try:
        doc=Document(str(path)); styles=getSampleStyleSheet()
        pdf_doc=SimpleDocTemplate(str(out),pagesize=A4,leftMargin=2*cm,rightMargin=2*cm,topMargin=2*cm,bottomMargin=2*cm)
        story=[]
        for para in doc.paragraphs:
            text=para.text.strip()
            if not text: story.append(Spacer(1,8)); continue
            sn="Heading1" if para.style.name.startswith("Heading 1") else "Heading2" if para.style.name.startswith("Heading 2") else "Heading3" if para.style.name.startswith("Heading") else "Normal"
            safe=text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(safe,styles[sn])); story.append(Spacer(1,4))
        for table in doc.tables:
            data=[[cell.text for cell in row.cells] for row in table.rows]
            if data:
                tbl=Table(data,repeatRows=1)
                tbl.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),rl_colors.HexColor("#d4380d")),("TEXTCOLOR",(0,0),(-1,0),rl_colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),9),("GRID",(0,0),(-1,-1),0.5,rl_colors.grey),("PADDING",(0,0),(-1,-1),4),("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.white,rl_colors.HexColor("#f7f4ef")])]))
                story.append(tbl); story.append(Spacer(1,12))
        if not story: story.append(Paragraph("(Empty document)",styles["Normal"]))
        pdf_doc.build(story)
    except Exception as e:
        path.unlink(missing_ok=True); raise HTTPException(500,f"Conversion failed: {e}")
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="converted.pdf", media_type="application/pdf")

@app.post("/api/pdf-to-excel")
async def pdf_to_excel(file: UploadFile=File(...)):
    path=save_upload(file); out=out_path("extracted.xlsx")
    try:
        wb=openpyxl.Workbook(); wb.remove(wb.active)
        hf=Font(bold=True,color="FFFFFF"); hfill=PatternFill("solid",fgColor="D4380D")
        afill=PatternFill("solid",fgColor="FFF1EC")
        with pdfplumber.open(str(path)) as pdf:
            found=0
            for pn,page in enumerate(pdf.pages,1):
                for ti,tbl in enumerate(page.extract_tables() or []):
                    found+=1; ws=wb.create_sheet(title=f"P{pn}_T{ti+1}"[:31])
                    for ri,row in enumerate(tbl):
                        for ci,val in enumerate(row):
                            cell=ws.cell(row=ri+1,column=ci+1,value=val or "")
                            if ri==0: cell.font=hf; cell.fill=hfill; cell.alignment=Alignment(horizontal="center")
                            elif ri%2==0: cell.fill=afill
                    for col in ws.columns:
                        ws.column_dimensions[col[0].column_letter].width=min(max((len(str(c.value or "")) for c in col),default=8)+4,50)
            if found==0:
                ws=wb.create_sheet(title="Text Content"); ws.append(["Page","Text"])
                ws["A1"].font=hf; ws["A1"].fill=hfill; ws["B1"].font=hf; ws["B1"].fill=hfill
                ws.column_dimensions["A"].width=8; ws.column_dimensions["B"].width=80
                with pdfplumber.open(str(path)) as pdf2:
                    for pn,page in enumerate(pdf2.pages,1):
                        for line in (page.extract_text() or "").split("\n"):
                            if line.strip(): ws.append([pn,line.strip()])
        wb.save(str(out))
    except Exception as e:
        path.unlink(missing_ok=True); raise HTTPException(500,f"Extraction failed: {e}")
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="extracted.xlsx",
                        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ── 18. PDF TO TEXT ────────────────────────────────────────────
@app.post("/api/pdf-to-text")
async def pdf_to_text(file: UploadFile=File(...)):
    path=save_upload(file); doc=fitz.open(str(path))
    all_text=[]
    for i,page in enumerate(doc):
        text=page.get_text("text")
        all_text.append(f"=== Page {i+1} ===\n{text}")
    doc.close(); out=out_path("extracted.txt")
    with open(out,"w",encoding="utf-8") as f: f.write("\n\n".join(all_text))
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="extracted.txt", media_type="text/plain")

# ── 19. ADD PAGE NUMBERS ───────────────────────────────────────
@app.post("/api/page-numbers")
async def add_page_numbers(file: UploadFile=File(...), position: str=Form("bottom-center"),
                            start_num: int=Form(1), font_size: int=Form(12)):
    path=save_upload(file); doc=fitz.open(str(path))
    for i,page in enumerate(doc):
        W,H=page.rect.width,page.rect.height; num=i+start_num
        pos={"bottom-center":(W/2-10,H-20),"bottom-right":(W-40,H-20),"bottom-left":(20,H-20),
             "top-center":(W/2-10,20),"top-right":(W-40,20),"top-left":(20,20)}
        x,y=pos.get(position,(W/2-10,H-20))
        page.insert_text((x,y),str(num),fontsize=font_size,color=(0.3,0.3,0.3),overlay=True)
    out=out_path("numbered.pdf"); doc.save(str(out)); doc.close(); path.unlink(missing_ok=True)
    return FileResponse(out, filename="numbered.pdf", media_type="application/pdf")

# ── 20. ADD WATERMARK ──────────────────────────────────────────
@app.post("/api/watermark")
async def add_watermark(file: UploadFile=File(...), watermark_text: str=Form("CONFIDENTIAL"),
                         opacity: float=Form(0.3), font_size: int=Form(48)):
    path=save_upload(file); doc=fitz.open(str(path))
    for page in doc:
        W,H=page.rect.width,page.rect.height
        page.insert_text((W/2-font_size*3,H/2),watermark_text,fontsize=font_size,
                         color=(0.7,0.7,0.7),overlay=False)
    out=out_path("watermarked.pdf"); doc.save(str(out)); doc.close(); path.unlink(missing_ok=True)
    return FileResponse(out, filename="watermarked.pdf", media_type="application/pdf")

# ── 21. EXTRACT PAGES ──────────────────────────────────────────
@app.post("/api/extract-pages")
async def extract_pages(file: UploadFile=File(...), page_range: str=Form(...)):
    path=save_upload(file); reader=PdfReader(str(path))
    indices=parse_page_range(page_range,len(reader.pages))
    if not indices: raise HTTPException(400,"Enter valid page numbers (e.g. 1-3, 5)")
    writer=PdfWriter()
    for i in indices: writer.add_page(reader.pages[i])
    out=out_path("extracted.pdf")
    with open(out,"wb") as f: writer.write(f)
    path.unlink(missing_ok=True)
    return FileResponse(out, filename="extracted_pages.pdf", media_type="application/pdf")

# ── 22. RESIZE PDF ─────────────────────────────────────────────
# @app.post("/api/resize")
# async def resize_pdf(file: UploadFile=File(...), page_size: str=Form("a4")):
#     size_map={"a4":fitz.paper_rect("a4"),"a3":fitz.paper_rect("a3"),
#               "letter":fitz.paper_rect("letter"),"legal":fitz.paper_rect("legal")}
#     target=size_map.get(page_size,fitz.paper_rect("a4"))
#     path=save_upload(file); src=fitz.open(str(path)); out_doc=fitz.open()
#     for page in src:
#         np=out_doc.new_page(width=target.width,height=target.height)
#         np.show_pdf_page(np.rect,src,page.number)
#     out=out_path("resized.pdf"); out_doc.save(str(out)); src.close(); path.unlink(missing_ok=True)
#     return FileResponse(out, filename="resized.pdf", media_type="application/pdf")

@app.post("/api/resize")
async def resize_pdf(file: UploadFile = File(...), page_size: str = Form("a4")):
    size_map = {
        "a4":     (595, 842),
        "a3":     (842, 1191),
        "letter": (612, 792),
        "legal":  (612, 1008),
    }
    w, h = size_map.get(page_size, (595, 842))
    path = save_upload(file)
    src = fitz.open(str(path))
    out_doc = fitz.open()
    for page in src:
        np = out_doc.new_page(width=w, height=h)
        # np.show_pdf_page(np.rect, src, page.number)
        np.show_pdf_page(np.rect, src, page.number, keep_proportion=False)
    out = out_path("resized.pdf")
    
    out_doc.save(str(out))
    out_doc.close()
    src.close()
    import gc; gc.collect()
    try:
       path.unlink(missing_ok=True)
    except:
       pass
    return FileResponse(out, filename="resized.pdf", media_type="application/pdf")

@app.post("/api/insert-image")
async def insert_image(
    file: UploadFile = File(...),
    image: UploadFile = File(...),
    page_num: int = Form(1),
    position: str = Form("center"),
    img_width: int = Form(200),
    img_height: int = Form(200),
):
    pdf_path = save_upload(file)
    img_path = save_upload(image)
    doc = fitz.open(str(pdf_path))
    page = doc[max(0, min(page_num - 1, len(doc) - 1))]
    W, H = page.rect.width, page.rect.height
    iw, ih = img_width, img_height
    pos_map = {
        "center":        ((W - iw) / 2,    (H - ih) / 2),
        "top-left":      (20,               20),
        "top-right":     (W - iw - 20,      20),
        "bottom-left":   (20,               H - ih - 20),
        "bottom-right":  (W - iw - 20,      H - ih - 20),
        "top-center":    ((W - iw) / 2,     20),
        "bottom-center": ((W - iw) / 2,     H - ih - 20),
    }
    x, y = pos_map.get(position, ((W - iw) / 2, (H - ih) / 2))
    rect = fitz.Rect(x, y, x + iw, y + ih)
    page.insert_image(rect, filename=str(img_path))
    out = out_path("image_inserted.pdf")
    doc.save(str(out))
    doc.close()
    pdf_path.unlink(missing_ok=True)
    img_path.unlink(missing_ok=True)
    return FileResponse(out, filename="image_inserted.pdf", media_type="application/pdf")

@app.post("/api/crop")
async def crop_pdf(
    file: UploadFile = File(...),
    top: int = Form(0),
    bottom: int = Form(0),
    left: int = Form(0),
    right: int = Form(0),
):
    pdf_path = save_upload(file)
    doc = fitz.open(str(pdf_path))
    for page in doc:
        r = page.rect
        new_rect = fitz.Rect(
            r.x0 + left,
            r.y0 + top,
            r.x1 - right,
            r.y1 - bottom
        )
        if new_rect.is_valid and new_rect.width > 0 and new_rect.height > 0:
            page.set_cropbox(new_rect)
    out = out_path("cropped.pdf")
    doc.save(str(out))
    doc.close()
    pdf_path.unlink(missing_ok=True)
    return FileResponse(out, filename="cropped.pdf", media_type="application/pdf")


@app.post("/api/batch-compress")
async def batch_compress(files: List[UploadFile] = File(...), level: str = Form("medium")):
    if len(files) < 2:
        raise HTTPException(400, "At least 2 PDF files required.")
    dpi = {"low": 150, "medium": 100, "high": 60}.get(level, 100)
    zip_out = out_path("batch_compressed.zip")
    with zipfile.ZipFile(zip_out, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in files:
            path = save_upload(file)
            doc = fitz.open(str(path))
            writer = fitz.open()
            for page in doc:
                mat = fitz.Matrix(dpi/72, dpi/72)
                pix = page.get_pixmap(matrix=mat)
                np = writer.new_page(width=page.rect.width, height=page.rect.height)
                np.insert_image(np.rect, pixmap=pix)
            buf = io.BytesIO()
            writer.save(buf, deflate=True, garbage=4)
            doc.close()
            writer.close()
            path.unlink(missing_ok=True)
            zf.writestr(f"compressed_{file.filename}", buf.getvalue())
    return FileResponse(zip_out, filename="batch_compressed.zip", media_type="application/zip")



@app.post("/api/batch-convert-to-pdf")
async def batch_convert_to_pdf(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, "Please upload at least one image.")
    zip_out = out_path("batch_pdfs.zip")
    with zipfile.ZipFile(zip_out, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in files:
            path = save_upload(file)
            img = Image.open(str(path)).convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="PDF")
            path.unlink(missing_ok=True)
            name = Path(file.filename).stem + ".pdf"
            zf.writestr(name, buf.getvalue())
    return FileResponse(zip_out, filename="batch_pdfs.zip", media_type="application/zip")



    
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    print(f"\n{'='*45}")
    print(f"  PDFBox Server Chal Raha Hai!")
    print(f"{'='*45}")
    print(f"  Browser mein kholein: http://localhost:{port}")
    print(f"  API Docs:             http://localhost:{port}/docs")
    print(f"  Band karna:           Ctrl+C")
    print(f"{'='*45}\n")
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=False)
